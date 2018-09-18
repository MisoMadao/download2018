from __future__ import print_function

import json

from scapy.all import *
from openpyxl import Workbook
from openpyxl.styles import Font, colors
from openpyxl.chart import PieChart, Reference, BarChart
from openpyxl.chart.series import DataPoint
from datetime import datetime
from docx import Document
import matplotlib.pyplot as plt
import numpy as np
from collections import OrderedDict
from graph_tool.all import *
import requests
import zipfile

"""
This class implements some enrichment and reporting we can do on traffic data
"""


class PAnalyzer:

    def __init__(self, fname='mycap.pcap', alexa_file=None, cu_file=None):
        """
        Here is your constructor, you can pass a pcap file name and also the csv for the alexa and cisco umbrella top
        1m domains
        """
        self.fname = fname
        self.packets = None
        self.alexa_file = alexa_file
        self.cu_file = cu_file
        self.read_pcap(self.fname)

    def capture(self):
        """capture traffic, it will be saved to packets attribute"""
        try:
            print('start capture...')
            self.packets = sniff()
        except KeyboardInterrupt:
            print('finish capture...')
            return

    def write_pcap(self):
        """save packets to file (constructor)"""
        wrpcap(self.fname, self.packets)

    def read_pcap(self, fname):
        """read pcap file (constructor) and save to packets attribute"""
        if os.path.exists(fname):
            self.packets = rdpcap(fname)
        else:
            print('file does not exist')

    def topSrcIP(self, top=10):
        """get the top <top> src ip"""
        src_ip = {}
        for pkt in self.packets:
            if IP in pkt:
                if pkt[IP].src not in src_ip:
                    src_ip[pkt[IP].src] = 0
                src_ip[pkt[IP].src] += 1
        src_ip = [(v, k) for k, v in src_ip.items()]
        src_ip.sort(reverse=True)
        return src_ip[:top]

    # can you adapt these methods in onw unique method?
    # can you make the new method so I can get top dst ip, top 10 port, top protocol by attribute passed to it?

    def topDstIP(self, top=10):
        """get the top <top> dst ip"""
        src_ip = {}
        for pkt in self.packets:
            if IP in pkt:
                if pkt[IP].dst not in src_ip:
                    src_ip[pkt[IP].dst] = 0
                src_ip[pkt[IP].dst] += 1
        src_ip = [(v, k) for k, v in src_ip.items()]
        src_ip.sort(reverse=True)
        return src_ip[:top]

    def excelReportIP(self, fname='ip_report.xlsx'):
        """create an excel with all packets and top IPs with graphs"""
        def set_headers(hrs, sheet, r=1):
            header_font = Font(color=colors.BLUE, bold=True)
            for h in hrs:
                cell = sheet.cell(row=r, column=hrs.index(h) + 1, value=h)
                cell.font = header_font
            return sheet

        if self.packets:
            wb = Workbook()
            ws = wb.active
            ws.title = 'packets'
            # headers
            headers = ['DateTime', 'Protocol', 'Source IP', 'Source Port', 'Destination IP', 'Destination Port']
            ws = set_headers(headers, ws)
            # values
            row = 2
            for pkt in self.packets:
                if IP in pkt:
                    ws.cell(row=row, column=headers.index('DateTime') + 1, value=datetime.fromtimestamp(pkt.time))
                    ws.cell(row=row, column=headers.index('Protocol') + 1, value=pkt[IP].proto)
                    ws.cell(row=row, column=headers.index('Source IP') + 1, value=pkt[IP].src)
                    ws.cell(row=row, column=headers.index('Destination IP') + 1, value=pkt[IP].dst)
                    if TCP in pkt[IP]:
                        layer = TCP
                    elif UDP in pkt[IP]:
                        layer = UDP
                    else:
                        layer = None
                    if layer:
                        ws.cell(row=row, column=headers.index('Source Port') + 1, value=pkt[IP][layer].sport)
                        ws.cell(row=row, column=headers.index('Destination Port') + 1, value=pkt[IP][layer].dport)
                    row += 1
            # charts
            wb.create_sheet('Charts')
            ws = wb['Charts']
            # top 10 src ip
            headers = ['Count', 'IP']
            ws.cell(row=1, column=1, value='Top 10 Source IP')
            ws = set_headers(headers, ws, r=2)
            row = 3
            for count, ip in self.topSrcIP():
                ws.cell(row=row, column=1, value=count)
                ws.cell(row=row, column=2, value=ip)
                row += 1
            # pie chart
            pie_chart = PieChart()
            labels = Reference(ws, min_col=2, min_row=3, max_row=row - 1)
            chart_data = Reference(ws, min_col=1, min_row=3, max_row=row - 1)
            pie_chart.add_data(chart_data, titles_from_data=True)
            pie_chart.set_categories(labels)
            pie_chart.title = 'Top 10 Source IP'

            # Cut the first slice out of the pie
            pie_slice = DataPoint(idx=0, explosion=20)
            pie_chart.series[0].data_points = [pie_slice]

            ws.add_chart(pie_chart, "F2")

            # top 10 dst ip
            headers = ['Count', 'IP']
            ws.cell(row=19, column=1, value='Top 10 Destination IP')
            ws = set_headers(headers, ws, r=20)
            row = 21
            for count, ip in self.topSrcIP():
                ws.cell(row=row, column=1, value=count)
                ws.cell(row=row, column=2, value=ip)
                row += 1
            # pie chart
            bar_chart = BarChart()
            bar_chart.type = "col"
            bar_chart.style = 10
            bar_chart.title = 'Top 10 Destination IP'
            bar_chart.y_axis.title = 'count'

            chart_data = Reference(ws, min_col=1, min_row=21, max_row=row - 1)
            cats = Reference(ws, min_col=2, min_row=21, max_row=row - 1)
            bar_chart.add_data(chart_data, titles_from_data=True)
            bar_chart.set_categories(cats)
            bar_chart.shape = 4
            ws.add_chart(bar_chart, "F20")

            wb.save(fname)
        else:
            print('no packets, no report')

    # the docx module doesn't have support for graphs currently. But it does have support for images.
    # can you use or adapt other methods of this class to create images for the top IPs and include them in the doc?

    def wordReportIP(self, fname='ip_report.docx'):
        """create a word document with all packets and top IPs"""
        if self.packets:
            document = Document()

            document.add_heading('Top 10 Dst IP', 0)
            for count, ip in self.topDstIP():
                document.add_paragraph('{}: {}'.format(ip, count), style='List Number')
            document.add_page_break()

            document.add_heading('Top 10 Src IP', 0)
            for count, ip in self.topSrcIP():
                document.add_paragraph('{}: {}'.format(ip, count), style='List Number')
            document.add_page_break()

            document.add_heading('Traffic', 0)
            p = document.add_paragraph('Packets Summary')
            p.bold = True
            p.italic = True
            # add table
            table = document.add_table(rows=1, cols=6)
            # headers
            headers = ['DateTime', 'Protocol', 'Source IP', 'Source Port', 'Destination IP', 'Destination Port']
            hdr_cells = table.rows[0].cells
            for h in headers:
                hdr_cells[headers.index(h)].text = h
            for pkt in self.packets:
                if IP in pkt:
                    row_cells = table.add_row().cells
                    row_cells[headers.index('DateTime')].text = '{}'.format(datetime.fromtimestamp(pkt.time))
                    row_cells[headers.index('Protocol')].text = '{}'.format(pkt[IP].proto)
                    row_cells[headers.index('Source IP')].text = '{}'.format(pkt[IP].src)
                    row_cells[headers.index('Destination IP')].text = '{}'.format(pkt[IP].dst)
                    if TCP in pkt[IP]:
                        layer = TCP
                    elif UDP in pkt[IP]:
                        layer = UDP
                    else:
                        layer = None
                    if layer:
                        row_cells[headers.index('Source Port')].text = '{}'.format(pkt[IP][layer].sport)
                        row_cells[headers.index('Destination Port')].text = '{}'.format(pkt[IP][layer].dport)

            document.save(fname)
        else:
            print('no packets, no report')

    def trafficReport(self, precision='min', size='B'):
        """this will plot an histogram with the traffic size over time"""
        time_values = ['min', 'sec', 'hour', 'day']
        size_values = ['B', 'KB', 'MB', 'GB']
        if precision not in time_values:
            print('precision should be any of {}. using default'.format(time_values))
            precision = 'min'

        # traffic in buckets of time
        buckets = {}
        for pkt in self.packets:
            dt = datetime.fromtimestamp(pkt.time)
            if precision == 'sec':
                bucket = '{}-{}-{}T{}:{}:{}'.format(dt.year, dt.month, dt.day, dt.hour, dt.minute, dt.second)
            elif precision == 'min':
                bucket = '{}-{}-{}T{}:{}'.format(dt.year, dt.month, dt.day, dt.hour, dt.minute)
            elif precision == 'hour':
                bucket = '{}-{}-{}T{}'.format(dt.year, dt.month, dt.day, dt.hour)
            else:
                bucket = '{}-{}-{}'.format(dt.year, dt.month, dt.day)
            if bucket not in buckets:
                buckets[bucket] = 0
            buckets[bucket] += len(pkt)

        # adjust the bytes size
        if size not in size_values:
            print('size should be any of {}. using default'.format(size_values))
            size = 'B'
        elif size != 'B':
            if size == 'KB':
                div = 1000
            elif size == 'MB':
                div = pow(1000, 2)
            else:
                div = pow(1000, 3)
            for bucket, val in buckets.items():
                buckets[bucket] = val / float(div)

        # sort the dict by keys
        b = OrderedDict(sorted(buckets.items()))

        # get the image
        x = np.array(range(0, len(b)))
        y = np.array([_ for _ in b.values()])
        my_xticks = b.keys()
        plt.xticks(x, my_xticks, rotation='vertical')
        plt.ylabel(size)
        plt.subplots_adjust(bottom=0.3)
        plt.plot(x, y)
        plt.savefig('traffic_report.png')

    def dnsAnswerReport(self, fname='dns_report.txt'):
        """get all dns requests and answers from the pcap in a file"""
        with open(fname, 'w') as f:
            for pkt in self.packets:
                if IP in pkt:
                    if UDP in pkt[IP]:
                        if DNS in pkt[IP][UDP]:
                            f.write('{}\t{}\n'.format(datetime.fromtimestamp(pkt.time), pkt[IP][UDP][DNSQR].qname))
                            for _ in range(pkt[IP][UDP][DNS].ancount):
                                f.write('{}\t\t{}\n'.format(datetime.fromtimestamp(pkt.time),
                                                            pkt[IP][UDP][DNSRR][_].rdata))

    def toGraph(self, fname='graph.png'):
        """print your traffic connections to a graph, every IP address is a node"""
        VERTEX = {}
        GRAPH = Graph()
        VERTEX_NAME = GRAPH.new_vertex_property("string")

        def add_vertex(ip):
            if ip not in VERTEX:
                VERTEX[ip] = {'vertex': None, 'connected_to': []}
                VERTEX[ip]['vertex'] = GRAPH.add_vertex()
                VERTEX_NAME[VERTEX[ip]['vertex']] = ip

        def add_edge(src, dst):
            if dst not in VERTEX[src]['connected_to']:
                VERTEX[src]['connected_to'].append(dst)
                VERTEX[dst]['connected_to'].append(src)
                GRAPH.add_edge(VERTEX[src]['vertex'], VERTEX[dst]['vertex'])

        for pkt in self.packets:
            if 'IP' in pkt:
                ip_src = pkt['IP'].src
                add_vertex(ip_src)
                ip_dst = pkt['IP'].dst
                add_vertex(ip_dst)
                add_edge(ip_src, ip_dst)

        graph_draw(GRAPH, vertex_text=VERTEX_NAME, vertex_font_size=18, output_size=(3000, 3000), output=fname)

    def _getTop1mFile(self, url, fname):
        """download zip file and extract it"""
        # get file
        try:
            r = requests.get(url, stream=True)
        except Exception as ex:
            print(ex)
            return
        # save file
        with open(fname, 'wb') as f:
            for chunk in r.iter_content(chunk_size=1024):
                f.write(chunk)
        # unzip file
        with zipfile.ZipFile(fname, 'r') as zip_ref:
            zip_ref.extractall(os.path.dirname(os.path.abspath(__file__)))

    def _downloadtop10Alexa(self, alexa='http://s3.amazonaws.com/alexa-static/top-1m-alexa.csv.zip'):
        """download alexa top1m zip"""
        self.alexa_file = 'alexa_top_1m.zip'
        if not os.path.exists(self.alexa_file):
            self._getTop1mFile(alexa, self.alexa_file)
            os.rename('top-1m-alexa.csv', 'top-1m-alexa.csv')

    def _downloadtop10CU(self, ciscoumbrella='http://s3-us-west-1.amazonaws.com/umbrella-static/top-1m-alexa.csv.zip'):
        """download cisco umbrella top1m zip"""
        self.cu_file = 'cu_top_1m.zip'
        if not os.path.exists(self.cu_file):
            self._getTop1mFile(ciscoumbrella, self.cu_file)
            os.rename('top-1m-alexa.csv', 'top-1m-cu.csv')

    def domainInAlexa(self, domain):
        """is the domain in the alexa top1m?"""
        if self.alexa_file:
            with open(self.alexa_file, 'r') as a:
                for line in a.readlines():
                    n, d = line.rstrip().split(',')
                    if d == domain:
                        return True
        return False

    def domainInCU(self, domain):
        """is the domain in the cisco umbrella top1m?"""
        if self.cu_file:
            with open(self.cu_file, 'r') as a:
                for line in a.readlines():
                    n, d = line.rstrip().split(',')
                    if d == domain:
                        return True
        return False

    # the following method does not consider the vt limitations (eg free has 4 requests per minute).
    # can you adapt it to consider it and wait with some timeouts?

    def ipVt(self, ip, apikey):
        """is the IP present on virus total?"""
        params = {
            'ip': ip,
            'apikey': apikey
        }
        r = requests.get('http://www.virustotal.com/vtapi/v2/ip-address/report', params=params)
        if r.status_code == 200:
            r = r.json()
            print(r.keys())
            if r['response_code'] == 1:
                print('{} has:\n'
                      '\towner {}\n'
                      '\tasn {}\n'
                      '\tdetected_referrer_samples {}\n'
                      '\tresolutions {}\n'
                      '\tdetected_communicating_samples {}\n'
                      '\tdetected_downloaded_samples {}'
                      ''.format(ip, r.setdefault('as_owner', ''), r.setdefault('asn', ''),
                                len(r.setdefault('detected_referrer_samples', '')), len(r.setdefault('resolutions', '')),
                                len(r.setdefault('detected_communicating_samples', '')),
                                len(r.setdefault('detected_downloaded_samples', ''))))
            else:
                print('response code {}'.format(r.json()['response_code']))
        else:
            print('request with status {}'.format(r.status_code))

    def toJson(self, p_count):
        """print a packet in json"""
        def pkt_to_json(pkt):
            json_pkt = {}
            max_levels = 3

            def loop_payloads(layer, fields, lev):
                if hasattr(layer, 'fields'):
                    for _ in layer.fields:
                        fields[_] = '{}'.format(layer.fields[_])
                if lev:
                    if hasattr(layer, 'payload'):
                        fields[layer._name] = loop_payloads(layer.payload, {}, lev - 1)
                return fields

            json_pkt = loop_payloads(pkt, json_pkt, max_levels)
            return json_pkt

        for p in self.packets:
            yield json.dumps(pkt_to_json(p))
            p_count -= 1
            if not p_count:
                break
